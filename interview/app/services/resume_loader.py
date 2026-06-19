"""
이력서 / 자기소개서 파일을 텍스트로 로드.

지원 포맷:
  - .txt   : 평문 텍스트 (utf-8 / utf-8-sig / cp949 / euc-kr 자동 감지)
  - .docx  : Microsoft Word        (python-docx 필요)
  - .pdf   : PDF                   (pypdf 필요)
  - .hwp   : 한글 5.0 바이너리      (olefile 필요)
  - .hwpx  : 한글 신포맷 (zip+XML)   (별도 의존성 없음)

사용:
  from app.services.resume_loader import load_resume
  text = load_resume("/path/to/resume.docx")
"""

import os
import struct
from typing import List


SUPPORTED_EXTS = (".txt", ".docx", ".pdf", ".hwp", ".hwpx")


def load_resume(path: str) -> str:
    """파일 경로를 받아 텍스트를 반환. 확장자로 자동 분기."""
    path = os.path.expanduser(path)
    if not os.path.exists(path):
        raise FileNotFoundError(f"파일을 찾을 수 없습니다: {path}")

    ext = os.path.splitext(path)[1].lower()
    loaders = {
        ".txt":  _load_txt,
        ".docx": _load_docx,
        ".pdf":  _load_pdf,
        ".hwp":  _load_hwp,
        ".hwpx": _load_hwpx,
    }
    if ext not in loaders:
        raise ValueError(
            f"지원하지 않는 형식: {ext}. "
            f"지원 형식: {SUPPORTED_EXTS}"
        )
    return loaders[ext](path).strip()


# ---------- 포맷별 구현 ----------

def _load_txt(path: str) -> str:
    """utf-8 우선, 한글 환경 대비 cp949/euc-kr 폴백."""
    for enc in ("utf-8", "utf-8-sig", "cp949", "euc-kr"):
        try:
            with open(path, encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    # 최후 수단: 깨진 문자는 무시
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


def _load_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError(
            "docx 처리에는 python-docx 가 필요합니다.\n"
            "  pip install python-docx"
        ) from e

    doc = Document(path)
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # 표 내용도 포함
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _load_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # 구버전 호환
        except ImportError as e:
            raise ImportError(
                "pdf 처리에는 pypdf 가 필요합니다.\n"
                "  pip install pypdf"
            ) from e

    reader = PdfReader(path)
    return "\n".join((page.extract_text() or "") for page in reader.pages)


# ---------- HWP (한글 5.0 바이너리) ----------

def _load_hwp(path: str) -> str:
    """OLE Compound + zlib + 본문 PARA_TEXT 레코드 추출."""
    try:
        import olefile
    except ImportError as e:
        raise ImportError(
            "hwp 처리에는 olefile 이 필요합니다.\n"
            "  pip install olefile"
        ) from e
    import zlib

    ole = olefile.OleFileIO(path)
    try:
        # 본문 섹션 목록
        sections = []
        for entry in ole.listdir():
            if entry and entry[0] == "BodyText" and len(entry) >= 2:
                name = entry[1]
                if name.startswith("Section"):
                    sections.append(int(name[len("Section"):]))
        sections.sort()
        if not sections:
            raise RuntimeError("본문 섹션을 찾지 못함. 정상 hwp 파일이 맞나요?")

        # 압축 여부 (FileHeader 37번째 바이트 LSB)
        header = ole.openstream("FileHeader").read()
        is_compressed = bool(header[36] & 0x01)

        out = []
        for s_idx in sections:
            data = ole.openstream(f"BodyText/Section{s_idx}").read()
            if is_compressed:
                data = zlib.decompress(data, -15)
            out.append(_extract_hwp_text(data))
        return "\n".join(out)
    finally:
        ole.close()


def _extract_hwp_text(section_data: bytes) -> str:
    """HWP 본문 바이너리에서 PARA_TEXT(=67) 레코드만 골라 텍스트로 변환."""
    HWPTAG_PARA_TEXT = 67  # HWPTAG_BEGIN(16) + 51
    parts: List[str] = []
    i = 0
    n = len(section_data)
    while i + 4 <= n:
        header_int = struct.unpack_from("<I", section_data, i)[0]
        rec_type = header_int & 0x3FF
        rec_len = (header_int >> 20) & 0xFFF
        if rec_len == 0xFFF:
            # 확장 길이 (다음 4바이트)
            if i + 8 > n:
                break
            rec_len = struct.unpack_from("<I", section_data, i + 4)[0]
            i += 8
        else:
            i += 4

        if rec_type == HWPTAG_PARA_TEXT:
            rec = section_data[i:i + rec_len]
            parts.append(_decode_hwp_para(rec))

        i += rec_len
    return "\n".join(p for p in parts if p)


def _decode_hwp_para(rec: bytes) -> str:
    """단락 레코드(UTF-16LE wchar 배열 + 인라인 컨트롤 문자) 디코드."""
    chars: List[str] = []
    j = 0
    n = len(rec)
    # 컨트롤 문자 중 14바이트 추가 데이터를 갖는 것들
    EXT_CTRLS = {1, 2, 3, 4, 11, 12, 14, 15, 16, 17, 18, 21}
    while j + 2 <= n:
        ch = struct.unpack_from("<H", rec, j)[0]
        j += 2
        if ch in (10, 13):
            chars.append("\n")
        elif ch < 32:
            if ch in EXT_CTRLS:
                j += 14   # 확장 컨트롤: 14바이트 스킵
            # 그 외 단순 컨트롤은 그냥 무시
        else:
            chars.append(chr(ch))
    return "".join(chars)


# ---------- HWPX (한글 신포맷, zip+xml) ----------

def _load_hwpx(path: str) -> str:
    """zip 안의 Contents/section*.xml 에서 텍스트 노드만 모음."""
    import zipfile
    from xml.etree import ElementTree as ET

    out: List[str] = []
    with zipfile.ZipFile(path) as z:
        names = sorted(
            n for n in z.namelist()
            if n.startswith("Contents/section") and n.endswith(".xml")
        )
        for name in names:
            xml_bytes = z.read(name)
            try:
                root = ET.fromstring(xml_bytes)
            except ET.ParseError:
                continue
            buf: List[str] = []
            for elem in root.iter():
                # 네임스페이스 무시하고 t 또는 text 태그만
                tag = elem.tag.split("}", 1)[-1]
                if tag in ("t", "text") and elem.text:
                    buf.append(elem.text)
            out.append("".join(buf))
    return "\n".join(out)


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("사용법: python resume_loader.py <file>")
        sys.exit(1)
    text = load_resume(sys.argv[1])
    print(f"--- 추출된 텍스트 ({len(text)}자) ---")
    print(text)
